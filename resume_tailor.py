#!/usr/bin/env python3
"""
Resume tailoring module with SEMANTIC REWRITING.
Uses AI to intelligently rewrite Professional Summary and Skills sections
based on job requirements.
"""

import os
import re
import shutil
from typing import Dict, List, Optional, Tuple
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pdfplumber
from config import Config
from logger import get_logger

# Initialize logger
logger = get_logger("resume_tailor")


class ResumeTailor:
    """
    Resume tailoring with semantic AI rewriting.
    
    Features:
    - Extracts and identifies resume sections (Summary, Skills, Experience)
    - Uses AI to semantically rewrite Professional Summary
    - Highlights keywords from job description in Skills section
    - Preserves original formatting while updating content
    """
    
    # Common section headers to identify resume sections
    SUMMARY_HEADERS = [
        'professional summary', 'summary', 'profile', 'objective',
        'career objective', 'about me', 'executive summary',
        'professional profile', 'career summary'
    ]
    
    SKILLS_HEADERS = [
        'skills', 'technical skills', 'core competencies', 'competencies',
        'key skills', 'areas of expertise', 'qualifications',
        'professional skills', 'expertise'
    ]
    
    EXPERIENCE_HEADERS = [
        'experience', 'work experience', 'professional experience',
        'employment history', 'work history', 'career history'
    ]
    
    def __init__(self):
        self.use_ai = Config.OPENAI_API_KEY and Config.OPENAI_API_KEY != ''
        self.client = None
        
        if self.use_ai:
            try:
                from openai import OpenAI
                self.client = OpenAI(api_key=Config.OPENAI_API_KEY)
                logger.info("Resume tailor initialized with AI semantic rewriting")
            except Exception as e:
                logger.error(f"OpenAI import failed: {e}")
                self.use_ai = False
        else:
            logger.info("Resume tailor initialized with keyword-based tailoring (no API key)")
    
    def tailor_resume(
        self, 
        original_resume_path: str, 
        job: Dict, 
        user_profile: Dict, 
        output_dir: str = "tailored_resumes"
    ) -> str:
        """
        Create a tailored version of the resume for a specific job.
        
        Args:
            original_resume_path: Path to the original resume
            job: Job dictionary with title, company, description
            user_profile: User profile with skills, experience, education
            output_dir: Directory to save tailored resumes
        
        Returns:
            Path to the tailored resume file
        """
        if not os.path.exists(original_resume_path):
            raise FileNotFoundError(f"Resume not found: {original_resume_path}")
        
        # Create output directory
        os.makedirs(output_dir, exist_ok=True)
        
        # Determine file type
        is_pdf = original_resume_path.lower().endswith('.pdf')
        is_docx = original_resume_path.lower().endswith('.docx')
        
        logger.info(f"Tailoring resume for {job.get('title')} at {job.get('company')}")
        
        if is_docx:
            return self._tailor_docx(original_resume_path, job, user_profile, output_dir)
        elif is_pdf:
            return self._tailor_pdf_to_docx(original_resume_path, job, user_profile, output_dir)
        else:
            raise ValueError("Unsupported resume format. Use PDF or DOCX.")
    
    def _extract_resume_sections(self, doc: Document) -> Dict[str, Tuple[int, int, str]]:
        """
        Extract and identify sections from a DOCX resume.
        
        Returns:
            Dictionary mapping section names to (start_index, end_index, content)
        """
        sections = {}
        current_section = None
        current_start = 0
        current_content = []
        
        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip().lower()
            
            # Check if this paragraph is a section header
            is_header = False
            
            # Check for Summary headers
            for header in self.SUMMARY_HEADERS:
                if header in text and len(text) < 50:
                    if current_section:
                        sections[current_section] = (
                            current_start, 
                            i - 1, 
                            '\n'.join(current_content)
                        )
                    current_section = 'summary'
                    current_start = i
                    current_content = []
                    is_header = True
                    break
            
            if not is_header:
                # Check for Skills headers
                for header in self.SKILLS_HEADERS:
                    if header in text and len(text) < 50:
                        if current_section:
                            sections[current_section] = (
                                current_start, 
                                i - 1, 
                                '\n'.join(current_content)
                            )
                        current_section = 'skills'
                        current_start = i
                        current_content = []
                        is_header = True
                        break
            
            if not is_header:
                # Check for Experience headers
                for header in self.EXPERIENCE_HEADERS:
                    if header in text and len(text) < 50:
                        if current_section:
                            sections[current_section] = (
                                current_start, 
                                i - 1, 
                                '\n'.join(current_content)
                            )
                        current_section = 'experience'
                        current_start = i
                        current_content = []
                        is_header = True
                        break
            
            # Add content to current section
            if not is_header and para.text.strip():
                current_content.append(para.text.strip())
        
        # Save last section
        if current_section:
            sections[current_section] = (
                current_start, 
                len(doc.paragraphs) - 1, 
                '\n'.join(current_content)
            )
        
        logger.debug(f"Found sections: {list(sections.keys())}")
        return sections
    
    def _generate_rewritten_summary(
        self, 
        original_summary: str, 
        job: Dict
    ) -> Optional[str]:
        """
        Use AI to semantically rewrite the Professional Summary.
        
        Args:
            original_summary: The current Professional Summary text
            job: Job dictionary with description, title, company
        
        Returns:
            AI-rewritten summary or None if AI unavailable
        """
        if not self.use_ai or not self.client:
            logger.warning("AI not available for semantic rewriting")
            return None
        
        job_title = job.get('title', 'the position')
        company = job.get('company', 'the company')
        job_description = job.get('description', '')[:2000]  # Limit size
        
        prompt = f"""You are an expert resume writer. Rewrite this Professional Summary to be highly relevant to the target job.

ORIGINAL SUMMARY:
{original_summary}

TARGET JOB:
- Title: {job_title}
- Company: {company}
- Description: {job_description}

INSTRUCTIONS:
1. Keep the SAME professional tone and length (2-4 sentences)
2. Emphasize skills and experiences that match the job description
3. Use keywords from the job description naturally
4. Highlight achievements relevant to this role
5. Keep formatting professional - no bullet points, just sentences
6. Do NOT add skills or experience the candidate doesn't have
7. Maintain first-person or third-person voice as in the original

Return ONLY the rewritten summary, nothing else."""

        try:
            logger.info("Generating AI-rewritten Professional Summary...")
            
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system", 
                        "content": "You are an expert resume writer who tailors professional summaries for specific job applications. Be concise and professional."
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.4,
                max_tokens=300
            )
            
            rewritten = response.choices[0].message.content.strip()
            
            # Clean up any quotes or extra formatting
            rewritten = rewritten.strip('"\'')
            
            logger.info("AI summary generation successful")
            logger.debug(f"Original: {original_summary[:100]}...")
            logger.debug(f"Rewritten: {rewritten[:100]}...")
            
            return rewritten
            
        except Exception as e:
            logger.error(f"AI summary generation failed: {e}")
            return None
    
    def _generate_enhanced_skills(
        self, 
        original_skills: str, 
        job: Dict
    ) -> Optional[str]:
        """
        Use AI to reorder and enhance the Skills section.
        Prioritizes skills mentioned in the job description.
        """
        if not self.use_ai or not self.client:
            return None
        
        job_description = job.get('description', '')[:1500]
        
        prompt = f"""Analyze these skills and the job description. Reorder the skills to prioritize those most relevant to the job.

CURRENT SKILLS SECTION:
{original_skills}

JOB DESCRIPTION:
{job_description}

INSTRUCTIONS:
1. Keep ALL existing skills - do not add new ones
2. Reorder to put most relevant skills FIRST
3. Maintain the original format (bullets, categories, etc.)
4. If the job mentions specific technologies, prioritize those
5. Keep it concise and well-organized

Return ONLY the reordered skills section, nothing else."""

        try:
            response = self.client.chat.completions.create(
                model="gpt-4",
                messages=[
                    {
                        "role": "system", 
                        "content": "You are an expert resume writer who optimizes skills sections for ATS systems."
                    },
                    {"role": "user", "content": prompt}
                ],
                temperature=0.3,
                max_tokens=400
            )
            
            return response.choices[0].message.content.strip()
            
        except Exception as e:
            logger.error(f"AI skills enhancement failed: {e}")
            return None
    
    def _tailor_docx(
        self, 
        original_path: str, 
        job: Dict, 
        user_profile: Dict, 
        output_dir: str
    ) -> str:
        """Tailor a DOCX resume with semantic rewriting."""
        
        if self.use_ai:
            return self._tailor_docx_with_ai(original_path, job, user_profile, output_dir)
        else:
            return self._tailor_docx_keyword_based(original_path, job, user_profile, output_dir)
    
    def _tailor_docx_with_ai(
        self, 
        original_path: str, 
        job: Dict, 
        user_profile: Dict, 
        output_dir: str
    ) -> str:
        """
        Use AI for SEMANTIC REWRITING of resume sections.
        
        This is the main enhancement - actually rewrites content based on job.
        """
        try:
            # Load document
            doc = Document(original_path)
            
            # Extract sections
            sections = self._extract_resume_sections(doc)
            
            if not sections:
                logger.warning("Could not identify resume sections, using keyword-based method")
                return self._tailor_docx_keyword_based(original_path, job, user_profile, output_dir)
            
            # Track what we've modified
            modifications = []
            
            # ========================================
            # 1. Rewrite Professional Summary
            # ========================================
            if 'summary' in sections:
                start_idx, end_idx, original_summary = sections['summary']
                
                # Generate rewritten summary
                new_summary = self._generate_rewritten_summary(original_summary, job)
                
                if new_summary:
                    # Find and replace the summary paragraphs
                    # We'll replace the first content paragraph after the header
                    summary_replaced = False
                    for i in range(start_idx + 1, min(end_idx + 1, len(doc.paragraphs))):
                        para = doc.paragraphs[i]
                        if para.text.strip() and not self._is_section_header(para.text):
                            # Replace this paragraph's content
                            # Preserve formatting of first run
                            if para.runs:
                                # Clear existing runs and add new text
                                for run in para.runs[1:]:
                                    run.text = ""
                                para.runs[0].text = new_summary
                            else:
                                para.text = new_summary
                            summary_replaced = True
                            modifications.append("Professional Summary")
                            break
                    
                    if not summary_replaced:
                        logger.warning("Could not locate summary paragraph to replace")
            
            # ========================================
            # 2. Optimize Skills Section (Reorder)
            # ========================================
            if 'skills' in sections:
                start_idx, end_idx, original_skills = sections['skills']
                
                # For skills, we'll highlight relevant ones by making them bold
                job_description = job.get('description', '').lower()
                
                for i in range(start_idx + 1, min(end_idx + 1, len(doc.paragraphs))):
                    para = doc.paragraphs[i]
                    if para.text.strip() and not self._is_section_header(para.text):
                        # Check if any word in this paragraph is in job description
                        for run in para.runs:
                            words = run.text.lower().split()
                            for word in words:
                                # Clean word
                                clean_word = re.sub(r'[^a-z0-9]', '', word)
                                if len(clean_word) > 3 and clean_word in job_description:
                                    run.bold = True
                                    break
                
                modifications.append("Skills (highlighted keywords)")
            
            # ========================================
            # 3. Generate output file
            # ========================================
            output_path = self._generate_output_path(job, output_dir, ".docx")
            
            # Add tailoring metadata at the very end (as a comment-like section)
            # This helps track which version goes with which job
            metadata_para = doc.add_paragraph()
            metadata_para.add_run("\n---\n").font.size = Pt(8)
            metadata_run = metadata_para.add_run(
                f"Tailored for: {job.get('title')} at {job.get('company')}"
            )
            metadata_run.font.size = Pt(8)
            metadata_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Save
            doc.save(output_path)
            
            logger.info(f"Semantic tailoring complete. Modified: {', '.join(modifications)}")
            logger.info(f"Saved to: {output_path}")
            
            return output_path
            
        except Exception as e:
            logger.error(f"AI tailoring failed: {e}")
            logger.info("Falling back to keyword-based tailoring")
            return self._tailor_docx_keyword_based(original_path, job, user_profile, output_dir)
    
    def _is_section_header(self, text: str) -> bool:
        """Check if text is likely a section header."""
        text_lower = text.strip().lower()
        all_headers = self.SUMMARY_HEADERS + self.SKILLS_HEADERS + self.EXPERIENCE_HEADERS
        return any(header in text_lower for header in all_headers) and len(text) < 50
    
    def _tailor_docx_keyword_based(
        self, 
        original_path: str, 
        job: Dict, 
        user_profile: Dict, 
        output_dir: str
    ) -> str:
        """Tailor DOCX resume by identifying and emphasizing relevant keywords."""
        doc = Document(original_path)
        
        # Extract job keywords
        job_description = (job.get('description', '') + ' ' + job.get('title', '')).lower()
        user_skills = [s.lower() for s in user_profile.get('skills', [])]
        
        # Find relevant skills mentioned in job description
        relevant_skills = [skill for skill in user_skills if skill in job_description]
        
        # Create a copy document
        tailored_doc = Document()
        
        # Add tailoring note at top
        note_para = tailored_doc.add_paragraph()
        note_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        note_para.add_run("TAILORED FOR: ").bold = True
        note_para.add_run(f"{job.get('title')} at {job.get('company')}")
        
        if relevant_skills:
            skills_para = tailored_doc.add_paragraph()
            skills_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            skills_para.add_run("RELEVANT SKILLS: ").bold = True
            skills_para.add_run(', '.join(relevant_skills[:10]))
        
        tailored_doc.add_paragraph()  # Spacing
        
        # Copy and enhance paragraphs
        for para in doc.paragraphs:
            new_para = tailored_doc.add_paragraph()
            new_para.alignment = para.alignment
            
            for run in para.runs:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                new_run.underline = run.underline
                
                # Bold relevant skills
                text_lower = run.text.lower()
                for skill in relevant_skills:
                    if skill in text_lower:
                        new_run.bold = True
                        break
        
        # Copy tables
        for table in doc.tables:
            new_table = tailored_doc.add_table(
                rows=len(table.rows), 
                cols=len(table.columns)
            )
            for i, row in enumerate(table.rows):
                for j, cell in enumerate(row.cells):
                    new_table.rows[i].cells[j].text = cell.text
        
        # Save
        output_path = self._generate_output_path(job, output_dir, ".docx")
        tailored_doc.save(output_path)
        
        logger.info(f"Keyword-based tailoring complete: {output_path}")
        return output_path
    
    def _tailor_pdf_to_docx(
        self, 
        original_path: str, 
        job: Dict, 
        user_profile: Dict, 
        output_dir: str
    ) -> str:
        """Convert PDF to DOCX and tailor it with AI if available."""
        # Extract text from PDF
        text = ""
        try:
            with pdfplumber.open(original_path) as pdf:
                for page in pdf.pages:
                    text += page.extract_text() or ""
        except Exception as e:
            logger.error(f"Error extracting PDF text: {e}")
            return self._copy_resume(original_path, job, output_dir)
        
        # Create new DOCX
        doc = Document()
        
        # If AI is available, try to rewrite the summary
        if self.use_ai:
            # Try to extract summary from text
            summary_text = self._extract_summary_from_text(text)
            if summary_text:
                rewritten_summary = self._generate_rewritten_summary(summary_text, job)
                if rewritten_summary:
                    # Add header
                    header = doc.add_paragraph()
                    header.add_run("PROFESSIONAL SUMMARY").bold = True
                    # Add rewritten summary
                    doc.add_paragraph(rewritten_summary)
                    doc.add_paragraph()  # Spacing
                    
                    # Add rest of resume (excluding old summary)
                    remaining_text = self._remove_summary_from_text(text, summary_text)
                    for line in remaining_text.split('\n'):
                        if line.strip():
                            doc.add_paragraph(line.strip())
                    
                    output_path = self._generate_output_path(job, output_dir, ".docx")
                    doc.save(output_path)
                    logger.info(f"PDF converted and tailored with AI: {output_path}")
                    return output_path
        
        # Fallback: basic conversion with tailoring note
        note_para = doc.add_paragraph()
        note_para.add_run("TAILORED FOR: ").bold = True
        note_para.add_run(f"{job.get('title')} at {job.get('company')}")
        doc.add_paragraph()
        
        for line in text.split('\n'):
            if line.strip():
                doc.add_paragraph(line.strip())
        
        output_path = self._generate_output_path(job, output_dir, ".docx")
        doc.save(output_path)
        logger.info(f"PDF converted with basic tailoring: {output_path}")
        return output_path
    
    def _extract_summary_from_text(self, text: str) -> Optional[str]:
        """Extract Professional Summary section from plain text."""
        lines = text.split('\n')
        in_summary = False
        summary_lines = []
        
        for i, line in enumerate(lines):
            line_lower = line.strip().lower()
            
            # Check if this line starts the summary section
            if any(header in line_lower for header in self.SUMMARY_HEADERS):
                in_summary = True
                continue
            
            # Check if we've hit a new section
            if in_summary:
                if any(header in line_lower for header in 
                       self.SKILLS_HEADERS + self.EXPERIENCE_HEADERS):
                    break
                if line.strip():
                    summary_lines.append(line.strip())
        
        if summary_lines:
            return ' '.join(summary_lines)
        return None
    
    def _remove_summary_from_text(self, text: str, summary: str) -> str:
        """Remove the summary section from text."""
        return text.replace(summary, '')
    
    def _generate_output_path(self, job: Dict, output_dir: str, extension: str) -> str:
        """Generate a safe output filename."""
        from datetime import datetime
        
        company_safe = "".join(
            c for c in job.get('company', 'Company') 
            if c.isalnum() or c in (' ', '-', '_')
        ).strip()[:30]
        
        title_safe = "".join(
            c for c in job.get('title', 'Position') 
            if c.isalnum() or c in (' ', '-', '_')
        ).strip()[:30]
        
        timestamp = datetime.now().strftime("%Y%m%d_%H%M")
        
        filename = f"{timestamp}_{company_safe}_{title_safe}_Resume{extension}"
        filename = filename.replace(' ', '_')
        
        return os.path.join(output_dir, filename)
    
    def _copy_resume(self, original_path: str, job: Dict, output_dir: str) -> str:
        """Fallback: just copy the resume with a new name."""
        ext = os.path.splitext(original_path)[1]
        output_path = self._generate_output_path(job, output_dir, ext)
        shutil.copy2(original_path, output_path)
        logger.info(f"Resume copied (no tailoring): {output_path}")
        return output_path


# Test the module if run directly
if __name__ == "__main__":
    print("Resume Tailor - Semantic Rewriting Module")
    print("=" * 50)
    
    tailor = ResumeTailor()
    print(f"AI Enabled: {tailor.use_ai}")
    
    if tailor.use_ai:
        print("\n✅ OpenAI connected - semantic rewriting available")
        print("   - Will rewrite Professional Summary based on job")
        print("   - Will optimize Skills section order")
    else:
        print("\n⚠️  No OpenAI API key - using keyword-based tailoring")
        print("   - Will highlight matching skills")
        print("   - Will add tailoring header")
    
    print("\nUsage:")
    print("  tailor.tailor_resume('resume.docx', job_dict, profile_dict)")
